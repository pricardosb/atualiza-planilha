import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import datetime
import calendar
import openpyxl
from openpyxl import load_workbook
from openpyxl.styles import Font
from copy import copy
import streamlit.components.v1 as components

def tentar_converter_numero(val):
    """Converte texto numérico em int/float nativo para o Excel reconhecer como número."""
    if pd.isna(val) or val == "" or val is None:
        return ""
    if isinstance(val, (int, float)):
        return val
    val_str = str(val).strip().replace(',', '.')
    try:
        num = float(val_str)
        return int(num) if num.is_integer() else num
    except (ValueError, TypeError):
        return str(val)

def limpar_texto_xml(texto):
    """Remove caracteres inválidos de controle ASCII que corrompem documentos Word (.docx)."""
    if pd.isna(texto) or texto is None:
        return ""
    texto_str = str(texto)
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', texto_str)


# =============================================================================
# --- FUNÇÕES DE EXPORTAÇÃO CORRIGIDAS ---
# =============================================================================

def gerar_excel_bytes(dados_exportacao):
    output = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Salvamento Remição"

    for item in dados_exportacao:
        ws.append([f"NOME: {item['nome']}"])
        ws.append([f"ORGANIZAÇÃO: {item['organiz']} | FUNÇÃO: {item['funcao']} | REMUNERAÇÃO: {item['remuneracao']} | SAÍDA: {item['saida']}"])
        ws.append([])

        pivot_df = item['pivot_df']
        if not pivot_df.empty:
            headers = ["ANO"] + list(pivot_df.columns)
            ws.append(headers)
            for idx_row, row_data in pivot_df.iterrows():
                # Converte dinamicamente para número real no Excel
                row_vals = [tentar_converter_numero(idx_row)] + [tentar_converter_numero(v) for v in row_data.values]
                ws.append(row_vals)

        ws.append(["Total de Dias:", tentar_converter_numero(item['total_dias'])])
        ws.append([])
        ws.append([])

    wb.save(output)
    return output.getvalue()


def gerar_docx_bytes(dados_exportacao):
    output = io.BytesIO()
    from docx import Document
    
    doc = Document()
    doc.add_heading("Espaço de Dados para Salvamento", level=1)

    for item in dados_exportacao:
        doc.add_heading(limpar_texto_xml(f"NOME: {item['nome']}"), level=2)
        p_meta = doc.add_paragraph()
        p_meta.add_run(
            limpar_texto_xml(
                f"ORGANIZAÇÃO: {item['organiz']} | "
                f"FUNÇÃO: {item['funcao']} | "
                f"REMUNERAÇÃO: {item['remuneracao']} | "
                f"SAÍDA: {item['saida']}"
            )
        )

        pivot_df = item['pivot_df']
        if not pivot_df.empty:
            headers = ["ANO"] + list(pivot_df.columns)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = limpar_texto_xml(h)

            for idx_row, row_data in pivot_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = limpar_texto_xml(idx_row)
                for i, val in enumerate(row_data.values):
                    row_cells[i+1].text = limpar_texto_xml(val)

        p_tot = doc.add_paragraph()
        p_tot.add_run(limpar_texto_xml(f"Total de Dias: {item['total_dias']}")).bold = True
        doc.add_paragraph()

    doc.save(output)
    return output.getvalue()

def extrair_mes_ano_do_nome(nome_arquivo):
    import re
    
    # Dicionário para converter o nome do mês escrito no arquivo em número
    meses = {
        "JANEIRO": "01", "FEVEREIRO": "02", "MARÇO": "03", "MARCO": "03",
        "ABRIL": "04", "MAIO": "05", "JUNHO": "06", "JULHO": "07",
        "AGOSTO": "08", "SETEMBRO": "09", "OUTUBRO": "10",
        "NOVEMBRO": "11", "DEZEMBRO": "12"
    }
    
    nome_upper = str(nome_arquivo).upper()
    
    # Procura um ano de 4 dígitos que comece com 20 (ex: 2023, 2024)
    ano_match = re.search(r'\b(20\d{2})\b', nome_upper)
    ano = ano_match.group(1) if ano_match else None
    
    # Procura o mês correspondente no nome do arquivo
    mes = None
    for nome_mes, num_mes in meses.items():
        if nome_mes in nome_upper:
            mes = num_mes
            break
            
    # Se achou mês e ano, retorna no formato MM/YYYY
    if mes and ano:
        return f"{mes}/{ano}"
    
    # Se falhar em achar um dos dois, retorna a mensagem padrão
    return "SEM MÊS/ANO"


# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="SINALE WEB", layout="wide")

# --- INICIALIZAÇÃO DE ESTADOS GLOBAIS ---
if "source_df" not in st.session_state:
    st.session_state["source_df"] = None
if "wb_data" not in st.session_state:
    st.session_state["wb_data"] = None
if "last_dest_name" not in st.session_state:
    st.session_state["last_dest_name"] = None
if "fila_modificacoes" not in st.session_state:
    st.session_state["fila_modificacoes"] = []
if "select_all" not in st.session_state:
    st.session_state["select_all"] = False
if "file_settings" not in st.session_state:
    st.session_state["file_settings"] = {}
if "pesquisa_df" not in st.session_state:
    st.session_state["pesquisa_df"] = None
if "executar_config" not in st.session_state:
    st.session_state["executar_config"] = False


# --- FUNÇÕES DE SUPORTE ---
def copiar_estilo_completo(origem, destino):
    if origem.has_style:
        destino.font = copy(origem.font)
        destino.border = copy(origem.border)
        destino.fill = copy(origem.fill)
        destino.number_format = copy(origem.number_format)
        destino.protection = copy(origem.protection)
        destino.alignment = copy(origem.alignment)


def deduplicar_colunas(colunas):
    vistos = {}
    novas_colunas = []
    for col in colunas:
        col_str = str(col).strip()
        if col_str in vistos:
            vistos[col_str] += 1
            novas_colunas.append(f"{col_str} ({vistos[col_str]})")
        else:
            vistos[col_str] = 1
            novas_colunas.append(col_str)
    return novas_colunas


def extrair_valor_limpo(df, idx, col_name):
    try:
        val = df.iloc[idx][col_name]
        if isinstance(val, pd.Series):
            val = val.iloc[0]
        if pd.isna(val):
            return None
        return val.item() if hasattr(val, 'item') else val
    except:
        return None


def converter_valor_inteligente(val_str, dtype_original):
    if val_str is None or str(val_str).strip() == "":
        return None
    val_str = str(val_str).strip()
    if pd.api.types.is_integer_dtype(dtype_original):
        try:
            return int(val_str)
        except ValueError:
            pass
    elif pd.api.types.is_float_dtype(dtype_original):
        try:
            return float(val_str.replace(',', '.'))
        except ValueError:
            pass
    try:
        return float(val_str.replace(',', '.'))
    except ValueError:
        return val_str


def formatar_datas_dataframe(df_input):
    df_out = df_input.copy()
    for col in df_out.columns:
        if pd.api.types.is_datetime64_any_dtype(df_out[col]):
            df_out[col] = df_out[col].dt.strftime('%d/%m/%Y').fillna('')
        else:
            df_out[col] = df_out[col].apply(
                lambda v: "" if pd.isna(v) else (
                    v.strftime('%d/%m/%Y') if isinstance(v, (datetime.datetime, datetime.date, pd.Timestamp))
                    else (str(v).split(' ')[0] if isinstance(v, str) and (' 00:00:00' in str(v) or 'T00:00:00' in str(v)) else v)
                )
            )
    return df_out


def calcular_pascoa(ano):
    a = ano % 19
    b = ano // 100
    c = ano % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    L = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * L) // 451
    mes = (h + L - 7 * m + 114) // 31
    dia = ((h + L - 7 * m + 114) % 31) + 1
    return datetime.date(ano, mes, dia)


def obter_estatisticas_mes(ano, mes):
    cal = calendar.monthcalendar(ano, mes)
    pascoa = calcular_pascoa(ano)
    feriados = [
        datetime.date(ano, 1, 1),
        pascoa - datetime.timedelta(days=47),
        pascoa - datetime.timedelta(days=2),
        datetime.date(ano, 4, 21),
        datetime.date(ano, 5, 1),
        pascoa + datetime.timedelta(days=60),
        datetime.date(ano, 9, 7),
        datetime.date(ano, 10, 12),
        datetime.date(ano, 11, 2),
        datetime.date(ano, 11, 15),
        datetime.date(ano, 11, 20),
        datetime.date(ano, 12, 25),
    ]
    feriados_mes = [f for f in feriados if f.month == mes and f.year == ano]

    dias_seg_sex_total = 0
    dias_seg_sab_total = 0
    feriados_seg_sex = 0
    feriados_seg_sab = 0
    lista_feriados_detalhes = []

    for semana in cal:
        for i in range(7):
            dia = semana[i]
            if dia != 0:
                data_atual = datetime.date(ano, mes, dia)
                wd = data_atual.weekday()
                if wd < 5:
                    dias_seg_sex_total += 1
                    dias_seg_sab_total += 1
                elif wd == 5:
                    dias_seg_sab_total += 1

                if data_atual in feriados_mes:
                    if wd < 5:
                        feriados_seg_sex += 1
                        feriados_seg_sab += 1
                        lista_feriados_detalhes.append((data_atual, "Seg a Sex"))
                    elif wd == 5:
                        feriados_seg_sab += 1
                        lista_feriados_detalhes.append((data_atual, "Sábado"))

    return {
        "seg_sex_total": dias_seg_sex_total,
        "seg_sex_feriados": feriados_seg_sex,
        "seg_sex_uteis": dias_seg_sex_total - feriados_seg_sex,
        "seg_sab_total": dias_seg_sab_total,
        "seg_sab_feriados": feriados_seg_sab,
        "seg_sab_uteis": dias_seg_sab_total - feriados_seg_sab,
        "feriados_detalhes": lista_feriados_detalhes
    }


def gerar_arquivo_atualizado_bytes(source_input, header, fila, df_original, sheet_name=None):
    wb = load_workbook(io.BytesIO(source_input) if isinstance(source_input, bytes) else source_input)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb[wb.sheetnames[0]]
    for mod in fila:
        col_target = mod['coluna']
        valor_convertido = converter_valor_inteligente(mod['novo_valor'], df_original[col_target].dtype)
        for idx in mod['indices']:
            excel_row = idx + header + 1
            ws.cell(row=excel_row, column=df_original.columns.get_loc(col_target) + 1, value=valor_convertido)

            if col_target.strip().upper() in ["SAIDA", "SAÍDA"]:
                for col_idx in range(1, ws.max_column + 1):
                    cell = ws.cell(row=excel_row, column=col_idx)
                    current_font = cell.font
                    if current_font:
                        cell.font = Font(
                            name=current_font.name,
                            size=current_font.size,
                            bold=current_font.bold,
                            italic=current_font.italic,
                            strike=current_font.strike,
                            underline=current_font.underline,
                            color="FF0000"
                        )
                    else:
                        cell.font = Font(color="FF0000")

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def titulo_estilizado(subtitulo=""):
    st.markdown(
        f"<div style='text-align: center; padding: 1.5rem; background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); color: white; border-radius: 12px; margin-bottom: 1.5rem;'><h1>⚡ SINALE WEB</h1><p>{subtitulo}</p></div>",
        unsafe_allow_html=True
    )


def obter_nome_coluna_por_letra(df, colunas_disponiveis, letra):
    mapa_letras = {
        'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4, 'F': 5, 'G': 6, 'H': 7,
        'I': 8, 'J': 9, 'K': 10, 'L': 11, 'M': 12, 'N': 13, 'O': 14,
        'P': 15, 'Q': 16, 'R': 17, 'S': 18, 'T': 19, 'U': 20, 'V': 21,
        'W': 22, 'X': 23, 'Y': 24, 'Z': 25
    }
    idx = mapa_letras.get(letra.upper())
    if idx is not None and idx < len(colunas_disponiveis):
        return colunas_disponiveis[idx]
    return None


def gerar_config_largura_colunas(df_subset, colunas):
    config = {}
    for col in colunas:
        if col in df_subset.columns:
            nome_coluna_upper = str(col).strip().upper()
            
            # 1. REGRA PARA A COLUNA "NOME": Tamanho baseado no CONTEÚDO
            if nome_coluna_upper == "NOME":
                tamanho_conteudo = df_subset[col].astype(str).str.len().max() if not df_subset[col].empty else 10
                if pd.isna(tamanho_conteudo):
                    tamanho_conteudo = 10
                
                # ~8 pixels por letra + margem
                largura_pixels = int(tamanho_conteudo * 8) + 20
                
                # Garante que não fique pequena demais nem ocupe a tela inteira sozinha
                largura_pixels = max(150, min(largura_pixels, 450))
                
            # 2. REGRA PARA AS DEMAIS COLUNAS: Tamanho baseado EXCLUSIVAMENTE no CABEÇALHO
            else:
                tamanho_titulo = len(str(col))
                largura_pixels = int(tamanho_titulo * 9) + 20
                largura_pixels = max(50, largura_pixels)
            
            config[col] = st.column_config.Column(width=largura_pixels)
            
    return config

# --- MENU PRINCIPAL ---
menu_opcao = st.sidebar.radio(
    "Selecione a rotina:",
    [
        "INCLUSÃO DE TRABALHO",
        "ATUALIZAÇÕES GERAIS",
        "PESQUISA PARA REMIÇÃO",
        "LIMPAR ARQUIVO",
        "SOMENTE TRABALHADORES ATIVOS",
        "SAIR DO SISTEMA"
    ]
)

# =============================================================================
# --- OPÇÃO 1: INCLUSÃO DE TRABALHO ---
# =============================================================================
if menu_opcao == "INCLUSÃO DE TRABALHO":
    titulo_estilizado("INTEGRADOR ==> DADOS GERAIS DO INTERNO >>> SINALE")
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("1. Arquivo de ORIGEM")
        source_file = st.file_uploader("Selecione o arquivo de ORIGEM", type=["xlsx", "xls", "csv", "txt"], key="src_upload")
        origem_tem_cabecalho = st.checkbox("Arquivo de Origem tem cabeçalho?", value=True)
    with col2:
        st.subheader("2. Arquivo de DESTINO")
        dest_file = st.file_uploader("Selecione o arquivo de DESTINO (.xlsx)", type=["xlsx"], key="dest_upload")
        header_dest = st.number_input("Linha do cabeçalho no Arquivo de Destino:", value=11, min_value=1)

    if source_file:
        cache_key_src = f"{source_file.name}_{origem_tem_cabecalho}"
        if "source_df" not in st.session_state or st.session_state.get("last_cache_key_src") != cache_key_src:
            hdr = 0 if origem_tem_cabecalho else None
            try:
                source_file.seek(0)
                ext = source_file.name.split('.')[-1].lower()
                engine_util = 'xlrd' if ext == 'xls' else ('openpyxl' if ext == 'xlsx' else None)
                raw = pd.read_excel(source_file, header=hdr, engine=engine_util)
                raw.columns = deduplicar_colunas(raw.columns) if origem_tem_cabecalho else [f"Col {i+1}" for i in range(len(raw.columns))]
                st.session_state["source_df"] = raw
                st.session_state["last_cache_key_src"] = cache_key_src
            except Exception as e:
                st.error(f"Erro ao ler arquivo: {e}")

    if dest_file:
        if "wb_data" not in st.session_state or st.session_state.get("last_dest_name") != dest_file.name:
            dest_file.seek(0)
            st.session_state["wb_data"] = dest_file.getvalue()
            st.session_state["last_dest_name"] = dest_file.name

    df_origem = st.session_state.get("source_df")
    wb_data = st.session_state.get("wb_data")

    if df_origem is not None and wb_data is not None:
        wb = load_workbook(io.BytesIO(wb_data))
        target_sheet = st.selectbox("Escolha a ABA na Planilha de Destino a ser Atualizada:", wb.sheetnames)
        ws = wb[target_sheet]

        st.subheader("3. Seleção de Registros")
        col_busca = st.selectbox("Coluna identificadora (para seleção):", df_origem.columns)
        opcoes_selecao = [f"{val} (Linha {idx})" for idx, val in df_origem[col_busca].items()]
        selected_options = st.multiselect("🔍 Escolha os registros:", opcoes_selecao)
        selected_indices = [int(item.split("(Linha ")[1].replace(")", "")) for item in selected_options]

        if selected_indices:
            st.info(f"📊 **{len(selected_indices)}** registro(s) selecionado(s) para atualização.")

        st.write("---")
        st.subheader("4. Correlação dos dados dos Arquivos ORIGEM X DESTINO")
        mapping = {}
        cols_ui = st.columns(4)
        opcoes_mapeamento = ["--- Não mapear ---", "⚠️ Auto-incrementar (Seq)"] + list(df_origem.columns)
        for i in range(1, ws.max_column + 1):
            header_val = ws.cell(row=header_dest, column=i).value
            with cols_ui[(i - 1) % 4]:
                map_val = st.selectbox(f"Col {i} ({header_val or 'S/ Título'})", opcoes_mapeamento, key=f"map_{i}")
                if map_val != "--- Não mapear ---":
                    mapping[i] = map_val

        st.write("---")
        st.subheader("5. Local da Atualização")
        modo_insercao = st.radio("Local de inserção:", ["Final da planilha", "A partir de uma linha específica"])
        target_row = st.number_input("Linha:", min_value=header_dest + 1, value=header_dest + 1) if modo_insercao == "A partir de uma linha específica" else ws.max_row + 1

        st.write("---")
        if st.button("🚀 Processar e Atualizar"):
            if not selected_indices:
                st.error("Selecione itens!")
                st.stop()
            ref_row_idx = (target_row - 1) if modo_insercao == "A partir de uma linha específica" else ws.max_row
            base_seq = 0
            if ref_row_idx >= header_dest:
                val_acima = ws.cell(row=ref_row_idx, column=1).value
                try:
                    base_seq = int(val_acima)
                except:
                    base_seq = 0
            if modo_insercao == "A partir de uma linha específica":
                ws.insert_rows(target_row, amount=len(selected_indices))
            current_row = target_row
            seq_val = base_seq
            for idx in selected_indices:
                seq_val += 1
                ref_row_idx = current_row - 1
                for col_idx in range(1, ws.max_column + 1):
                    target_cell = ws.cell(row=current_row, column=col_idx)
                    ref_cell = ws.cell(row=ref_row_idx, column=col_idx)
                    copiar_estilo_completo(ref_cell, target_cell)
                    if col_idx == 1 or mapping.get(col_idx) == "⚠️ Auto-incrementar (Seq)":
                        target_cell.value = seq_val
                    elif col_idx in mapping:
                        target_cell.value = extrair_valor_limpo(df_origem, idx, mapping[col_idx])
                    else:
                        target_cell.value = ref_cell.value
                current_row += 1
            buffer = io.BytesIO()
            wb.save(buffer)
            st.session_state["wb_data"] = buffer.getvalue()
            st.success("✅ Processamento concluído com sucesso!")
            st.download_button(
                "📥 Baixar Versão Atualizada",
                st.session_state["wb_data"],
                "sinale_atualizado.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

# =============================================================================
# --- OPÇÃO 2: ATUALIZAÇÕES GERAIS ---
# =============================================================================
elif menu_opcao == "ATUALIZAÇÕES GERAIS":
    titulo_estilizado("Atualizações Gerais")

    if st.session_state.get("wb_data") is not None:
        st.info("📁 Arquivo carregado automaticamente da memória.")
        if st.checkbox("🗑️ Descartar dados da memória e carregar novo arquivo", value=False, key="desc_op2"):
            st.session_state["wb_data"] = None
            st.session_state["fila_modificacoes"] = []
            st.success("Memória limpa com sucesso!")
            st.rerun()
    else:
        st.warning("⚠️ Nenhum arquivo de destino encontrado na memória. Faça o upload abaixo.")
        sinale_file = st.file_uploader("Selecione o arquivo do SINALE (.xlsx)", type=["xlsx"], key="upload_op2")
        if sinale_file:
            st.session_state["wb_data"] = sinale_file.getvalue()
            st.session_state["last_sinale_name"] = sinale_file.name
            st.rerun()

    if st.session_state.get("wb_data") is not None:
        wb_temp = load_workbook(io.BytesIO(st.session_state["wb_data"]), data_only=True)
        target_sheet = st.selectbox("Escolha a ABA do arquivo para trabalhar:", wb_temp.sheetnames, key="aba_op2")
        header = st.number_input("Linha do cabeçalho:", value=11, min_value=1, key="header_op2")
        df = pd.read_excel(io.BytesIO(st.session_state["wb_data"]), sheet_name=target_sheet, header=header - 1)

        st.subheader("🔍 Filtros de Visualização")
        cols_para_ver = st.multiselect("Quais campos deseja visualizar?", df.columns.tolist(), default=df.columns.tolist())
        col_filtro, val_filtro = st.columns(2)
        with col_filtro:
            filtro_col = st.selectbox("Coluna para buscar:", df.columns, key="filtro_col_op2")
        valores_existentes = sorted([str(v) for v in df[filtro_col].dropna().unique()])
        with val_filtro:
            filtro_vals = st.multiselect("Selecione o(s) valor(es) para filtrar:", valores_existentes, key="filtro_vals_op2")

        df_view = df.copy()
        if filtro_vals:
            df_view = df_view[df_view[filtro_col].astype(str).isin(filtro_vals)]
        st.metric("Total de Registros Encontrados", len(df_view))

        df_view_fmt = formatar_datas_dataframe(df_view[cols_para_ver])
        st.dataframe(df_view_fmt, use_container_width=True, hide_index=True)

        st.subheader("✏️ Seleção para Atualizar")
        if "select_all" not in st.session_state:
            st.session_state["select_all"] = False
        cols_btns = st.columns([1, 1, 4])
        with cols_btns[0]:
            if st.button("✅ Marcar Todos", key="btn_marcar_t"):
                st.session_state["select_all"] = True
                st.rerun()
        with cols_btns[1]:
            if st.button("❌ Desmarcar Todos", key="btn_desmarcar_t"):
                st.session_state["select_all"] = False
                st.rerun()

        df_for_edit = df_view.copy()
        df_for_edit.insert(0, "Atualizar?", st.session_state["select_all"])
        df_editado = st.data_editor(
            df_for_edit,
            column_config={"Atualizar?": st.column_config.CheckboxColumn()},
            use_container_width=True,
            key="editor_op2"
        )

        selecionados = df_editado[df_editado["Atualizar?"] == True]
        st.metric("Total de Registros Marcados", len(selecionados))

        if not selecionados.empty:
            col_target = st.selectbox("Selecione a coluna que deseja alterar:", df.columns, key="col_target_op2")
            if col_target.strip().upper() == "DIAS":
                st.markdown("---")
                st.subheader("📅 Cálculo Automático de Dias Úteis (Seg a Sáb / Seg a Sex)")
                c_mes, c_ano = st.columns(2)
                meses_dict = {
                    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
                    "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
                }
                with c_mes:
                    mes_escolhido_nome = st.selectbox("Selecione o Mês:", list(meses_dict.keys()), key="sel_mes_dias")
                    mes_num = meses_dict[mes_escolhido_nome]
                with c_ano:
                    ano_escolhido = st.number_input("Digite o Ano:", min_value=2020, max_value=2035, value=datetime.date.today().year, key="sel_ano_dias")
                stats = obter_estatisticas_mes(ano_escolhido, mes_num)
                st.info(f"**Resumo para {mes_escolhido_nome}/{ano_escolhido}:**\n* **Segunda a Sábado:** {stats['seg_sab_total']} brutos | **Úteis:** **{stats['seg_sab_uteis']}**\n* **Segunda a Sexta:** {stats['seg_sex_total']} brutos | **Úteis:** **{stats['seg_sex_uteis']}**")

            valores_antigos_str = ", ".join([str(v) for v in selecionados[col_target].dropna().unique()])
            st.info(f"📌 **Valor(es) atual(is) / antigo(s)** no campo **'{col_target}'**: **{valores_antigos_str if valores_antigos_str else 'Vazio'}**")
            novo_val = st.text_input("Digite o novo valor:", key="novo_val_op2")

            if st.button("➕ Adicionar à Fila de Modificações", key="btn_add_fila"):
                st.session_state["fila_modificacoes"].append({
                    "indices": selecionados.index.tolist(),
                    "coluna": col_target,
                    "novo_valor": novo_val,
                    "valor_antigo": valores_antigos_str,
                    "vl_busca": ", ".join(filtro_vals) if filtro_vals else "Todos",
                    "aba": target_sheet
                })
                st.success("Modificação adicionada à fila!")
                st.rerun()

        if st.session_state["fila_modificacoes"]:
            st.markdown("---")
            st.subheader("📋 Fila de Modificações Pendentes")
            df_fila_resumo = pd.DataFrame([
                {
                    "Remover?": False,
                    "ID_ITEM": i,
                    "ABA": item.get("aba", "Geral"),
                    "CAMPO": item.get("coluna", ""),
                    "NOVO VALOR": item.get("novo_valor", "")
                }
                for i, item in enumerate(st.session_state["fila_modificacoes"])
            ])
            df_fila_editado = st.data_editor(
                df_fila_resumo,
                column_config={"Remover?": st.column_config.CheckboxColumn("Remover?"), "ID_ITEM": None},
                disabled=["ABA", "CAMPO", "NOVO VALOR"],
                use_container_width=True,
                key="editor_fila"
            )
            col_f1, col_f2, col_f3 = st.columns(3)
            with col_f1:
                if st.button("🗑️ Remover Selecionados"):
                    indices = df_fila_editado[df_fila_editado["Remover?"] == True]["ID_ITEM"].tolist()
                    st.session_state["fila_modificacoes"] = [
                        item for i, item in enumerate(st.session_state["fila_modificacoes"]) if i not in indices
                    ]
                    st.rerun()
            with col_f3:
                file_bytes = gerar_arquivo_atualizado_bytes(
                    io.BytesIO(st.session_state["wb_data"]),
                    header,
                    st.session_state["fila_modificacoes"],
                    df,
                    sheet_name=target_sheet
                )
                st.download_button(
                    "📥 Baixar Arquivo Atualizado",
                    file_bytes,
                    "sinale_atualizado_final.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )



# =============================================================================
# --- OPÇÃO 3: PESQUISA PARA REMIÇÃO ---
# =============================================================================
elif menu_opcao == "PESQUISA PARA REMIÇÃO":
    titulo_estilizado("Pesquisa para Remição")

    # Inicializa a chave dinâmica para zerar os componentes visuais no "Limpar Tudo"
    if "uploader_key" not in st.session_state:
        st.session_state["uploader_key"] = 0

    st.subheader("1. Configuração de Arquivos, Abas e Campos")
    
    col_btn_1, col_btn_2 = st.columns(2)
    with col_btn_1:
        uploaded_files = st.file_uploader(
            "1. Selecione os arquivos (.xlsx, .xls, .ods)",
            type=["xlsx", "xls", "ods"],
            accept_multiple_files=True,
            key=f"search_upload_{st.session_state['uploader_key']}"
        )
    
    qtd_arquivos = len(uploaded_files) if uploaded_files else 0
    st.info(f"📊 **Quantidade de arquivos selecionados:** {qtd_arquivos}")

    with col_btn_2:
        st.markdown("<br>", unsafe_allow_html=True)
        fazer_upload_btn = st.button("2. Fazer Upload e Configurar Abas", key="btn_fazer_upload_op3", type="primary")

    if fazer_upload_btn:
        if uploaded_files:
            st.session_state["executar_config"] = True
            st.session_state["rolar_apos_upload"] = True
            st.success("Arquivos carregados com sucesso! Configure as abas abaixo:")
        else:
            st.error("Selecione pelo menos um arquivo antes de fazer o upload.")
            st.session_state["executar_config"] = False
            st.session_state["rolar_apos_upload"] = False

    if uploaded_files and st.session_state.get("executar_config"):
        settings = {}
        for f_idx, f in enumerate(uploaded_files):
            file_key = f"{f_idx}_{f.name}"
            f_bytes = f.getvalue()
            file_ext = f.name.split('.')[-1].lower()
            
            try:
                engine_val = 'odf' if file_ext == 'ods' else None
                xl = pd.ExcelFile(io.BytesIO(f_bytes), engine=engine_val)
                sheets_available = xl.sheet_names
            except Exception as e:
                st.error(f"Erro ao ler o arquivo {f.name}: {e}.")
                continue

            pref_sheets = [s for s in sheets_available if any(p in s.strip().upper() for p in ["COM REMUNER", "SEM REMUNER", "DEM_COM", "DEM_SEM"])]

            if pref_sheets:
                default_sheets = pref_sheets
                is_fallback = False
            else:
                default_sheets = [sheets_available[0]] if sheets_available else []
                is_fallback = True

            with st.expander(f"📁 Configurações para: Arquivo {f_idx+1} - {f.name}", expanded=True):
                selected_sheets = st.multiselect(
                    f"Selecione aba(s) para {f.name}",
                    sheets_available,
                    default=default_sheets,
                    key=f"sheets_{file_key}_{st.session_state['uploader_key']}"
                )

                sheet_config = {}
                for i, sheet in enumerate(selected_sheets):
                    st.markdown(f"**Aba: `{sheet}`**")

                    sheet_upper = sheet.strip().upper()
                    if "DEM_COM" in sheet_upper:
                        default_header = 17
                    elif "DEM_SEM" in sheet_upper:
                        default_header = 19
                    elif any(p in sheet_upper for p in ["COM REMUNER", "SEM REMUNER"]):
                        default_header = 11
                    else:
                        default_header = 10 if is_fallback else 11

                    header_row = st.number_input(
                        f"Linha do cabeçalho para aba '{sheet}'",
                        value=default_header,
                        min_value=1,
                        key=f"head_{file_key}_{sheet}_{st.session_state['uploader_key']}"
                    )

                    try:
                        df_preview = pd.read_excel(io.BytesIO(f_bytes), sheet_name=sheet, header=header_row - 1, nrows=0, engine=engine_val)
                        cols_aba = [str(c).strip() for c in df_preview.columns]
                    except:
                        cols_aba = []

                    default_col = None
                    for c in cols_aba:
                        c_up = str(c).strip().upper()
                        if c_up in ["NOME DO INTERNO", "NOME DO INTERNO "]:
                            default_col = c
                            break
                    if not default_col:
                        for c in cols_aba:
                            if str(c).strip().upper() == "NOME":
                                default_col = c
                                break
                    if not default_col:
                        for c in cols_aba:
                            if str(c).strip().upper().startswith("NOME"):
                                default_col = c
                                break
                    if not default_col:
                        for c in cols_aba:
                            if "NOME" in str(c).strip().upper():
                                default_col = c
                                break
                    if not default_col and len(cols_aba) > 8:
                        default_col = cols_aba[8]
                    elif not default_col and cols_aba:
                        default_col = cols_aba[0]

                    opcoes_colunas = ["--- Não pesquisar nesta aba ---"] + cols_aba
                    default_idx = opcoes_colunas.index(default_col) if default_col in opcoes_colunas else 0

                    col_escolhida = st.selectbox(
                        f"Selecione o campo (coluna) para a pesquisa na aba '{sheet}':",
                        opcoes_colunas,
                        index=default_idx,
                        key=f"col_search_{file_key}_{sheet}_{st.session_state['uploader_key']}"
                    )

                    sheet_config[sheet] = {
                        "header_idx": header_row - 1,
                        "col_busca": col_escolhida if col_escolhida != "--- Não pesquisar nesta aba ---" else None
                    }
                    st.markdown("---")

                settings[file_key] = sheet_config

        btn_consolidar = st.button("🔍 Carregar e Consolidar Dados para Pesquisa", key="btn_consolidar_op3", type="primary")

        # ROLAGEM AUTOMÁTICA ATÉ O FINAL
        if st.session_state.get("rolar_apos_upload"):
            components.html(
                """
                <script>
                    function rolarAteOFinal() {
                        const doc = window.parent.document;
                        const container = doc.querySelector('section.main') || doc.querySelector('[data-testid="stMain"]') || doc.querySelector('[data-testid="stAppViewContainer"]') || doc.documentElement;
                        if (container) {
                            container.scrollTo({ top: container.scrollHeight, behavior: 'smooth' });
                        }
                    }
                    setTimeout(rolarAteOFinal, 400);
                    setTimeout(rolarAteOFinal, 800);
                </script>
                """,
                height=0
            )
            st.session_state["rolar_apos_upload"] = False

        if btn_consolidar:
            all_results = []
            for f_idx, f in enumerate(uploaded_files):
                file_key = f"{f_idx}_{f.name}"
                f_bytes = f.getvalue()
                file_ext = f.name.split('.')[-1].lower()
                engine_val = 'odf' if file_ext == 'ods' else None
                
                try:
                    xl = pd.ExcelFile(io.BytesIO(f_bytes), engine=engine_val)
                    mes_ano_arquivo = extrair_mes_ano_do_nome(f.name)
                except:
                    mes_ano_arquivo = "SEM MÊS/ANO"

                file_cfg = settings.get(file_key, {})
                for sheet, cfg in file_cfg.items():
                    try:
                        df_tmp = pd.read_excel(io.BytesIO(f_bytes), sheet_name=sheet, header=cfg["header_idx"], engine=engine_val)
                        df_tmp.columns = [str(c).strip() for c in df_tmp.columns]
                        df_tmp.columns = deduplicar_colunas(df_tmp.columns)

                        col_pedida = cfg.get("col_busca")
                        target_col = None
                        if col_pedida:
                            for c in df_tmp.columns:
                                if str(c).strip().upper() == str(col_pedida).strip().upper():
                                    target_col = c
                                    break
                        if not target_col:
                            for c in df_tmp.columns:
                                if "NOME DO INTERNO" in str(c).strip().upper():
                                    target_col = c
                                    break
                        if not target_col:
                            for c in df_tmp.columns:
                                if "NOME" in str(c).strip().upper():
                                    target_col = c
                                    break
                        if not target_col and len(df_tmp.columns) > 8:
                            target_col = df_tmp.columns[8]
                        elif not target_col and len(df_tmp.columns) > 0:
                            target_col = df_tmp.columns[0]

                        if target_col and target_col in df_tmp.columns:
                            colunas_originais = list(df_tmp.columns)

                            df_tmp["Aba Original"] = sheet
                            df_tmp["Campo Pesquisado"] = target_col

                            val_nome = df_tmp[target_col].astype(str).str.strip()
                            df_tmp["Nome (Visualização)"] = val_nome
                            df_tmp["NOME_LIMPO"] = val_nome.str.upper()

                            df_tmp = df_tmp[~df_tmp["NOME_LIMPO"].isin(['', 'NAN', 'NONE', '0', 'NAT', 'NC', 'N/C'])].copy()

                            aba_upper = sheet.strip().upper()
                            is_dem_com = "DEM_COM" in aba_upper
                            is_dem_sem = "DEM_SEM" in aba_upper
                            is_com_remuner = "COM REMUNER" in aba_upper
                            is_sem_remuner = "SEM REMUNER" in aba_upper
                            col_f = obter_nome_coluna_por_letra(df_tmp, colunas_originais, 'F')
                            
                            usar_padrao_antigo = False
                            usar_dem_sem_antigo = False

                            is_03_a_05_2023 = False
                            is_06_a_07_2023 = False
                            is_08_2023 = False

                            if mes_ano_arquivo != "SEM MÊS/ANO":
                                try:
                                    mes_str, ano_str = mes_ano_arquivo.split('/')
                                    mes_val, ano_val = int(mes_str), int(ano_str)
                                    
                                    if ano_val == 2023 and mes_val in [3, 4, 5]:
                                        is_03_a_05_2023 = True
                                    elif ano_val == 2023 and mes_val in [6, 7]:
                                        is_06_a_07_2023 = True
                                    elif ano_val == 2023 and mes_val == 8:
                                        is_08_2023 = True

                                    if ano_val < 2025 or (ano_val == 2025 and mes_val < 9):
                                        usar_padrao_antigo = True

                                    if ano_val < 2019 or (ano_val == 2019 and mes_val < 11):
                                        usar_dem_sem_antigo = True
                                except Exception:
                                    pass

                            def extrair_dados_e_categoria(row):
                                if is_03_a_05_2023:
                                    if is_dem_com or is_com_remuner:
                                        cat = "COM REMUNERAÇÃO"
                                        letras = ["I", "B", "T", "V", "W", "X", "Y"]
                                    elif is_dem_sem or is_sem_remuner:
                                        cat = "SEM REMUNERAÇÃO"
                                        letras = ["J", "B", "S", "U", "V", "W", "X"]
                                    else:
                                        val_f = row[col_f] if (col_f and col_f in row) else None
                                        is_sim = str(val_f).strip().upper() == "SIM" if pd.notna(val_f) else False
                                        cat = "COM REMUNERAÇÃO" if is_sim else "SEM REMUNERAÇÃO"
                                        letras = ["I", "B", "T", "V", "W", "X", "Y"] if is_sim else ["J", "B", "S", "U", "V", "W", "X"]
                                        
                                elif is_06_a_07_2023:
                                    if is_dem_com or is_com_remuner:
                                        cat = "COM REMUNERAÇÃO"
                                        letras = ["I", "B", "U", "W", "X", "Y", "Z"]
                                    elif is_dem_sem or is_sem_remuner:
                                        cat = "SEM REMUNERAÇÃO"
                                        letras = ["J", "B", "S", "V", "W", "X", "Y"]
                                    else:
                                        val_f = row[col_f] if (col_f and col_f in row) else None
                                        is_sim = str(val_f).strip().upper() == "SIM" if pd.notna(val_f) else False
                                        cat = "COM REMUNERAÇÃO" if is_sim else "SEM REMUNERAÇÃO"
                                        letras = ["I", "B", "U", "W", "X", "Y", "Z"] if is_sim else ["J", "B", "S", "V", "W", "X", "Y"]

                                elif is_08_2023:
                                    if is_dem_com or is_com_remuner:
                                        cat = "COM REMUNERAÇÃO"
                                        letras = ["I", "B", "R", "T", "U", "V", "W"]
                                    elif is_dem_sem or is_sem_remuner:
                                        cat = "SEM REMUNERAÇÃO"
                                        letras = ["I", "B", "Q", "S", "T", "U", "V"]
                                    else:
                                        val_f = row[col_f] if (col_f and col_f in row) else None
                                        is_sim = str(val_f).strip().upper() == "SIM" if pd.notna(val_f) else False
                                        cat = "COM REMUNERAÇÃO" if is_sim else "SEM REMUNERAÇÃO"
                                        letras = ["I", "B", "R", "T", "U", "V", "W"] if is_sim else ["I", "B", "Q", "S", "T", "U", "V"]

                                else:
                                    if is_dem_com:
                                        cat = "COM REMUNERAÇÃO"
                                        letras = ["I", "B", None, "S", "T", "U", "V"]

                                    elif is_dem_sem:
                                        cat = "SEM REMUNERAÇÃO"
                                        if usar_dem_sem_antigo:
                                            letras = ["I", "B", "Y", "R", "S", "T", "U"]
                                        else:
                                            letras = ["I", "B", "Y", "S", "T", "U", "V"]

                                    elif is_com_remuner:
                                        cat = "COM REMUNERAÇÃO"
                                        if usar_padrao_antigo:
                                            letras = ["I", "B", "Q", "S", "T", "U", "V"]
                                        else:
                                            letras = ["B", "I", "J", "T", "U", "V", "W"]

                                    elif is_sem_remuner:
                                        cat = "SEM REMUNERAÇÃO"
                                        letras = ["I", "B", "W", "R", "S", "T", "U"]

                                    else:
                                        val_f = row[col_f] if (col_f and col_f in row) else None
                                        is_sim = str(val_f).strip().upper() == "SIM" if pd.notna(val_f) else False

                                        cat = "COM REMUNERAÇÃO" if is_sim else "SEM REMUNERAÇÃO"
                                        letras = ["J", "C", "X", "S", "T", "U", "V"]

                                row_vals = {
                                    "Categoria_Aba": cat,
                                    "LABEL_EXIBICAO": f"{mes_ano_arquivo} - {cat}"
                                }
                                for idx_p, let in enumerate(letras):
                                    if let is None:
                                        val = ""
                                        header_title = ""
                                    else:
                                        col_n = obter_nome_coluna_por_letra(df_tmp, colunas_originais, let)
                                        val = row[col_n] if col_n and col_n in row else None
                                        header_title = str(col_n) if col_n else f"Campo {idx_p+1}"
                                    row_vals[f"POS_{idx_p}"] = val
                                    row_vals[f"HEADER_{idx_p}"] = header_title

                                return pd.Series(row_vals)

                            res_df = df_tmp.apply(extrair_dados_e_categoria, axis=1)
                            df_tmp["MÊS/ANO - ABA"] = res_df["LABEL_EXIBICAO"]

                            df_processed = pd.concat([
                                df_tmp[[
                                    "MÊS/ANO - ABA",
                                    "Aba Original",
                                    "Campo Pesquisado",
                                    "Nome (Visualização)",
                                    "NOME_LIMPO"
                                ]],
                                res_df
                            ], axis=1)
                            all_results.append(df_processed)
                    except Exception as e:
                        st.error(f"Erro ao ler {f.name} - Aba {sheet}: {e}")

            if all_results:
                st.session_state["pesquisa_df"] = pd.concat(all_results, ignore_index=True)
                st.success(f"Dados consolidados com sucesso! **{len(st.session_state['pesquisa_df'])}** registros carregados.")
            else:
                st.warning("Nenhum dado encontrado com as configurações informadas.")
                st.session_state["pesquisa_df"] = None

    if st.session_state.get("pesquisa_df") is not None:
        df_pesq = st.session_state["pesquisa_df"]
        st.markdown("---")
        st.subheader("🔍 Filtros de Visualização e Busca")

        col_ord1, col_ord2 = st.columns([2, 2])
        with col_ord1:
            ordem_escolhida = st.radio(
                "📅 Ordenação por Mês/Ano:",
                ["Crescente (Antigo ➔ Recente)", "Decrescente (Recente ➔ Antigo)"],
                horizontal=True,
                key=f"ordem_radio_{st.session_state['uploader_key']}"
            )
        
        is_ascending = True if "Crescente" in ordem_escolhida else False

        nomes_disponiveis = sorted(df_pesq["Nome (Visualização)"].dropna().unique())
        nomes_selecionados = st.multiselect(
            "🔍 Digite para pesquisar e selecione o(s) nome(s):",
            options=nomes_disponiveis,
            key=f"busca_nomes_{st.session_state['uploader_key']}"
        )

        df_view = df_pesq.copy()
        if nomes_selecionados:
            df_view = df_view[df_view["Nome (Visualização)"].isin(nomes_selecionados)]

        st.metric("Total de Registros Encontrados", len(df_view))

        if not df_view.empty:
            
            def extrair_chave_data(val):
                try:
                    data_str = str(val).split(' - ')[0].strip()
                    if data_str == "SEM MÊS/ANO":
                        return 999999 if is_ascending else -1
                    m, y = data_str.split('/')
                    return int(y) * 100 + int(m)
                except:
                    return 999999 if is_ascending else -1
            
            df_view['chave_ordenacao'] = df_view['MÊS/ANO - ABA'].apply(extrair_chave_data)
            df_view = df_view.sort_values(by=['chave_ordenacao'], ascending=is_ascending).drop(columns=['chave_ordenacao'])

            df_display_all = formatar_datas_dataframe(df_view)

            def formatar_sem_decimal(val):
                if pd.isna(val) or str(val).strip() in ["", "nan", "None"]:
                    return ""
                try:
                    num = float(val)
                    return str(int(round(num)))
                except (ValueError, TypeError):
                    return str(val).strip()

            def conv_num(val):
                try:
                    v_str = str(val).replace(',', '.').strip()
                    return float(v_str) if v_str not in ["", "nan", "None"] else 0.0
                except:
                    return 0.0

            grupos_categorias = [
                ("🟢 COM REMUNERAÇÃO", "COM REMUNERAÇÃO", "com_rem"),
                ("🟡 SEM REMUNERAÇÃO", "SEM REMUNERAÇÃO", "sem_rem")
            ]

            mapa_meses = {
                "01": "JAN", "1": "JAN", "02": "FEV", "2": "FEV",
                "03": "MAR", "3": "MAR", "04": "ABR", "4": "ABR",
                "05": "MAI", "5": "MAI", "06": "JUN", "6": "JUN",
                "07": "JUL", "7": "JUL", "08": "AGO", "8": "AGO",
                "09": "SET", "9": "SET", "10": "OUT", "11": "NOV", "12": "DEZ"
            }
            ordem_meses_siglas = ["JAN", "FEV", "MAR", "ABR", "MAI", "JUN", "JUL", "AGO", "SET", "OUT", "NOV", "DEZ"]

            for titulo_grupo, cat_key, prefixo_key in grupos_categorias:
                df_grupo = df_display_all[df_display_all["Categoria_Aba"] == cat_key]

                if not df_grupo.empty:
                    pos_cols = [c for c in df_grupo.columns if str(c).startswith("POS_")]
                    pos_cols.sort(key=lambda x: int(x.split("_")[1]))

                    cabecalhos_padrao = ["NOME", "ORGANIZ", "FUNÇÃO", "ENTRADA", "SAIDA", "PREV", "REAL"]
                    rename_map = {}
                    
                    for idx_p, pos_col in enumerate(pos_cols):
                        if idx_p < len(cabecalhos_padrao):
                            rename_map[pos_col] = cabecalhos_padrao[idx_p]
                        else:
                            rename_map[pos_col] = f"Campo {idx_p+1}"

                    cols_exibir = ["MÊS/ANO - ABA"] + pos_cols
                    df_render = df_grupo[cols_exibir].rename(columns=rename_map)
                    df_render = df_render.rename(columns={"MÊS/ANO - ABA": "MES/ANO - ABA"})

                    if "REAL" in df_render.columns:
                        df_render["REAL"] = df_render["REAL"].apply(formatar_sem_decimal)

                    st.markdown(f"### {titulo_grupo} ({len(df_render)} registro(s))")

                    key_select = f"select_all_{prefixo_key}"
                    if key_select not in st.session_state:
                        st.session_state[key_select] = False

                    col_b1, col_b2, _ = st.columns([1, 1, 4])
                    with col_b1:
                        if st.button("✅ Marcar Todos", key=f"btn_marcar_{prefixo_key}_{st.session_state['uploader_key']}"):
                            st.session_state[key_select] = True
                            st.rerun()
                    with col_b2:
                        if st.button("❌ Desmarcar Todos", key=f"btn_desmarcar_{prefixo_key}_{st.session_state['uploader_key']}"):
                            st.session_state[key_select] = False
                            st.rerun()

                    df_render.insert(0, "SELECIONAR?", st.session_state[key_select])

                    col_config_conteudo = gerar_config_largura_colunas(df_render, df_render.columns.tolist())
                    col_config_conteudo["SELECIONAR?"] = st.column_config.CheckboxColumn("SELECIONAR?", default=False)

                    df_editado_res = st.data_editor(
                        df_render,
                        column_config=col_config_conteudo,
                        use_container_width=True,
                        hide_index=True,
                        key=f"editor_res_{prefixo_key}_{st.session_state['uploader_key']}"
                    )

                    selecionados_grupo = df_editado_res[df_editado_res["SELECIONAR?"] == True]
                    
                    if not selecionados_grupo.empty:
                        st.markdown("---")
                        st.markdown(f"### 💾 Espaço de Dados para Salvamento — {titulo_grupo}")
                        
                        nomes_unicos = selecionados_grupo["NOME"].dropna().unique()
                        dados_exportacao_grupo = []

                        for nome_interno in nomes_unicos:
                            df_nome_sel = selecionados_grupo[selecionados_grupo["NOME"] == nome_interno]
                            
                            organiz_val = ", ".join([str(v) for v in df_nome_sel["ORGANIZ"].dropna().unique() if str(v).strip() != ""])
                            funcao_val = ", ".join([str(v) for v in df_nome_sel["FUNÇÃO"].dropna().unique() if str(v).strip() != ""])
                            saida_val = ", ".join([str(v) for v in df_nome_sel["SAIDA"].dropna().unique() if str(v).strip() != ""])
                            
                            soma_real = sum(conv_num(v) for v in df_nome_sel["REAL"])
                            total_dias_nome = int(round(soma_real))

                            st.markdown(
                                f"**NOME:** {nome_interno} &nbsp;|&nbsp; "
                                f"**ORGANIZAÇÃO:** {organiz_val if organiz_val else 'N/A'} &nbsp;|&nbsp; "
                                f"**FUNÇÃO:** {funcao_val if funcao_val else 'N/A'} &nbsp;|&nbsp; "
                                f"**REMUNERAÇÃO:** {cat_key} &nbsp;|&nbsp; "
                                f"**SAÍDA:** {saida_val if saida_val else 'N/A'}"
                            )
                            
                            matrix_data = []
                            for _, r_row in df_nome_sel.iterrows():
                                raw_mes_ano_aba = str(r_row.get("MES/ANO - ABA", ""))
                                data_mes_ano = raw_mes_ano_aba.split(" - ")[0] if " - " in raw_mes_ano_aba else raw_mes_ano_aba
                                
                                mes_sigla, ano_str = "N/A", "N/A"
                                if "/" in data_mes_ano:
                                    parts = data_mes_ano.split("/")
                                    if len(parts) == 2:
                                        mes_num, ano_str = parts[0].strip(), parts[1].strip()
                                        mes_sigla = mapa_meses.get(mes_num, mes_num)
                                
                                val_real = formatar_sem_decimal(r_row.get("REAL", ""))
                                
                                matrix_data.append({
                                    "ANO": ano_str,
                                    "MÊS": mes_sigla,
                                    "REAL": val_real
                                })
                            
                            df_pivot = pd.DataFrame()
                            if matrix_data:
                                df_mat = pd.DataFrame(matrix_data)
                                
                                df_pivot = df_mat.pivot_table(
                                    index="ANO",
                                    columns="MÊS",
                                    values="REAL",
                                    aggfunc=lambda x: " / ".join([str(v) for v in x if pd.notna(v) and str(v).strip() != ""])
                                ).fillna("")
                                
                                cols_meses = sorted(df_pivot.columns, key=lambda m: ordem_meses_siglas.index(m) if m in ordem_meses_siglas else 99)
                                df_pivot = df_pivot[cols_meses]
                                
                                st.dataframe(df_pivot, use_container_width=True)
                            
                            st.markdown(f"**Total de Dias:** {total_dias_nome}")
                            st.markdown("<br>", unsafe_allow_html=True)

                            dados_exportacao_grupo.append({
                                "nome": nome_interno,
                                "organiz": organiz_val if organiz_val else "N/A",
                                "funcao": funcao_val if funcao_val else "N/A",
                                "remuneracao": cat_key,
                                "saida": saida_val if saida_val else "N/A",
                                "pivot_df": df_pivot,
                                "total_dias": total_dias_nome
                            })

                        if dados_exportacao_grupo:
                            st.markdown("#### 📥 Baixar Relatório do Salvamento")
                            col_dl1, col_dl2 = st.columns(2)
                            
                            excel_bytes = gerar_excel_bytes(dados_exportacao_grupo)
                            with col_dl1:
                                st.download_button(
                                    label="📊 Baixar Excel (.xlsx)",
                                    data=excel_bytes,
                                    file_name=f"salvamento_remicao_{prefixo_key}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    key=f"btn_dl_xlsx_{prefixo_key}_{st.session_state['uploader_key']}"
                                )
                            
                            docx_bytes = gerar_docx_bytes(dados_exportacao_grupo)
                            with col_dl2:
                                st.download_button(
                                    label="📄 Baixar Word (.docx)",
                                    data=docx_bytes,
                                    file_name=f"salvamento_remicao_{prefixo_key}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"btn_dl_docx_{prefixo_key}_{st.session_state['uploader_key']}"
                                )

                        total_marcados = len(selecionados_grupo)
                        st.caption(f"📌 **{total_marcados}** item(ns) selecionado(s) nesta tabela.")
                        st.markdown("---")
        else:
            st.info("ℹ️ Nenhum registro selecionado ou encontrado na pesquisa.")
            
    # =========================================================================
    # BOTÃO LIMPAR TUDO - RESET COMPLETO DE TELA E COMPONENTES
    # =========================================================================
    if st.button("🗑️ Limpar Tudo", key="btn_limpar_tudo_op3"):
        chave_atual = st.session_state.get("uploader_key", 0) + 1
        st.session_state.clear()
        st.session_state["uploader_key"] = chave_atual
        st.rerun()

